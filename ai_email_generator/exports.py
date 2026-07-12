"""Export helpers for generated emails."""

from __future__ import annotations

import re
from io import BytesIO
from typing import Iterable
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .models import GeneratedEmail


def safe_filename(value: str, *, suffix: str = "txt") -> str:
    """Create a filesystem-friendly filename."""

    slug = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_").lower()
    if not slug:
        slug = "ai_email"
    return f"{slug}.{suffix.lstrip('.')}"


def _build_text(email: GeneratedEmail) -> str:
    return email.full_text.strip() + "\n"


def export_email_txt(email: GeneratedEmail) -> bytes:
    return _build_text(email).encode("utf-8")


def export_email_md(email: GeneratedEmail) -> bytes:
    markdown = f"# {email.subject.strip()}\n\n{email.body.strip()}\n"
    return markdown.encode("utf-8")


def export_email_docx(email: GeneratedEmail) -> bytes:
    buffer = BytesIO()
    document = Document()

    title = document.add_heading("AI Generated Email", level=1)
    title.runs[0].font.size = Pt(18)

    subject = document.add_paragraph()
    subject.add_run("Subject: ").bold = True
    subject.add_run(email.subject.strip())

    document.add_paragraph("")

    for paragraph in email.body.strip().split("\n\n"):
        document.add_paragraph(paragraph.strip())

    document.add_paragraph("")
    footer = document.add_paragraph()
    footer.add_run(f"Model: {email.model}").italic = True

    document.save(buffer)
    return buffer.getvalue()


def export_email_pdf(email: GeneratedEmail) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
    )

    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="EmailSubject",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#14213d"),
            alignment=TA_LEFT,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EmailBody",
            parent=styles["BodyText"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=15,
            textColor=colors.HexColor("#1f2937"),
            alignment=TA_LEFT,
            spaceAfter=9,
        )
    )
    styles.add(
        ParagraphStyle(
            name="EmailMeta",
            parent=styles["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=9,
            leading=11,
            textColor=colors.HexColor("#6b7280"),
            alignment=TA_LEFT,
        )
    )

    story = [
        Paragraph(escape(email.subject.strip()), styles["EmailSubject"]),
        Spacer(1, 0.12 * inch),
    ]

    for paragraph in email.body.strip().split("\n\n"):
        escaped = escape(paragraph.strip()).replace("\n", "<br/>")
        story.append(Paragraph(escaped, styles["EmailBody"]))

    story.append(Spacer(1, 0.12 * inch))
    story.append(Paragraph(escape(f"Model: {email.model}"), styles["EmailMeta"]))

    doc.build(story)
    return buffer.getvalue()
